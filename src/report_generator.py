"""
标准化科研图表报告生成模块
============================

基于 python-docx 一键生成包含交互式时序图表、
相关性统计表的完整 Word 科研报告。
"""

import os
import datetime
from typing import Dict, List, Optional

import numpy as np
import pandas as pd

from .config_manager import ConfigManager


class ReportGenerator:
    """科研报告生成器"""

    def __init__(self, config: Optional[ConfigManager] = None):
        self.config = config or ConfigManager()

    def _try_import_docx(self):
        """尝试导入 python-docx 库"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            return Document, Inches, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, qn
        except ImportError:
            raise ImportError(
                "请安装 python-docx 库: pip install python-docx"
            )

    def _setup_chinese_font(self, run, font_name: str = "宋体", size: int = 12, bold: bool = False):
        """设置中文字体"""
        _, _, Pt, _, RGBColor, _, _, qn = self._try_import_docx()
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    def _add_heading_cn(self, doc, text: str, level: int = 1):
        """添加中文标题"""
        _, _, Pt, _, _, WD_ALIGN_PARAGRAPH, _, qn = self._try_import_docx()
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            self._setup_chinese_font(run, size=18 - level * 2, bold=True)
        return heading

    def _add_paragraph_cn(self, doc, text: str, font_size: int = 12,
                        bold: bool = False, alignment: str = "left"):
        """添加中文段落"""
        _, _, _, _, _, WD_ALIGN_PARAGRAPH, _, _ = self._try_import_docx()
        para = doc.add_paragraph()
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(text)
        self._setup_chinese_font(run, size=font_size, bold=bold)
        return para

    def _add_dataframe_table(self, doc, df: pd.DataFrame, title: str = "",
                            float_digits: int = 4):
        """将DataFrame转换为Word表格"""
        _, _, _, _, _, _, WD_TABLE_ALIGNMENT, _ = self._try_import_docx()

        if title:
            self._add_paragraph_cn(doc, title, font_size=12, bold=True)

        n_rows, n_cols = df.shape
        table = doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col_name)
            for para in cell.paragraphs:
                for run in para.runs:
                    self._setup_chinese_font(run, size=10, bold=True)

        for i in range(n_rows):
            for j in range(n_cols):
                cell = table.rows[i + 1].cells[j]
                value = df.iloc[i, j]
                if isinstance(value, float):
                    cell.text = f"{value:.{float_digits}f}"
                else:
                    cell.text = str(value)
                for para in cell.paragraphs:
                    for run in para.runs:
                        self._setup_chinese_font(run, size=10)

        doc.add_paragraph()

    def _add_image_safe(self, doc, image_path: str, width_inches: float = 6.0):
        """安全添加图片"""
        _, Inches, _, _, _, _, _, _ = self._try_import_docx()
        if image_path and os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(width_inches))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = 1
            except Exception as e:
                self._add_paragraph_cn(doc, f"[图片插入失败: {image_path} - {str(e)}]")
        else:
            self._add_paragraph_cn(doc, f"[图片文件不存在: {image_path}]")

    def generate_research_report(self,
                                timeseries_df: pd.DataFrame,
                                corr_report: Dict,
                                cleaning_report: pd.DataFrame,
                                figure_paths: Dict[str, Dict],
                                chronology_df: Optional[pd.DataFrame] = None,
                                core_id: Optional[str] = None,
                                output_filename: Optional[str] = None) -> str:
        """
        生成完整的科研报告

        Args:
            timeseries_df: 标准化时序数据
            corr_report: 相关性分析报告
            cleaning_report: 数据清洗报告
            figure_paths: 图表文件路径
            chronology_df: 年代学数据
            core_id: 钻孔ID
            output_filename: 输出文件名

        Returns:
            生成的Word报告文件路径
        """
        Document, Inches, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, qn = self._try_import_docx()

        if core_id is None:
            core_id = self.config.paths.get("default_core", "unknown")
        if output_filename is None:
            output_filename = f"{core_id}_research_report_{datetime.datetime.now().strftime('%Y%m%d')}.docx"

        output_path = self.config.get_report_path(output_filename)

        doc = Document()

        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

        self._add_heading_cn(doc, "湿地泥炭沉积多指标古气候分析报告", level=0)
        self._add_paragraph_cn(doc, f"钻孔编号: {core_id}", font_size=12, alignment="center")
        self._add_paragraph_cn(doc, f"报告生成日期: {datetime.datetime.now().strftime('%Y年%m月%d日')}", font_size=12, alignment="center")
        self._add_paragraph_cn(doc, f"编制单位: 湿地古气候研究所", font_size=12, alignment="center")

        doc.add_page_break()

        self._add_heading_cn(doc, "一、研究概述", level=1)
        self._add_paragraph_cn(doc,
            "本报告基于泥炭沉积多代用指标（碳同位素、植物残体、重金属）数据，"
            "通过年代-深度模型重建了研究区万年尺度的水文和温度演化历史。"
            "分析流程包括原始数据自动清洗、年代学插值校准、多指标皮尔逊相关性分析以及"
            "多层联动时序可视化。",
            font_size=12, alignment="justify"
        )

        try:
            data_source = self.config.get_data_source(core_id)
            self._add_paragraph_cn(doc, f"研究地点: {data_source.get('location', '未知')}", font_size=12)
            self._add_paragraph_cn(doc, f"数据描述: {data_source.get('description', '')}", font_size=12)
        except Exception:
            pass

        self._add_heading_cn(doc, "二、数据质量与清洗", level=1)
        self._add_paragraph_cn(doc,
            "原始实验数据经过严格的质量控制，主要清洗步骤包括：仪器噪声平滑（Savitzky-Golay滤波）、"
            "异常离群值剔除（IQR方法，1.5倍四分位距）、超出地球化学合理范围的数据过滤、"
            "以及缺失值处理。各代用指标的清洗统计如下：",
            font_size=12, alignment="justify"
        )

        cleaning_display = cleaning_report.copy()
        cleaning_display.columns = [
            "代用指标", "原始数据点数", "噪声平滑处理", "离群值剔除",
            "超出范围剔除", "缺失值剔除", "最终有效点数"
        ]
        self._add_dataframe_table(doc, cleaning_display, "表1 各代用指标数据清洗统计")

        self._add_heading_cn(doc, "三、年代学框架", level=1)
        if chronology_df is not None:
            self._add_paragraph_cn(doc,
                f"本次研究共获得 {len(chronology_df)} 个可靠的测年控制点，"
                f"年代范围覆盖 {chronology_df['age_yrBP'].min():.0f} – {chronology_df['age_yrBP'].max():.0f} yr BP，"
                f"对应沉积深度 {chronology_df['depth_cm'].min():.1f} – {chronology_df['depth_cm'].max():.1f} cm。"
                f"采用线性插值方法构建连续的年代-深度转换模型，时间分辨率为10年。",
                font_size=12, alignment="justify"
            )

            ad_png = figure_paths.get("age_depth", {}).get("png", "")
            if ad_png:
                self._add_image_safe(doc, ad_png, width_inches=5.5)
                self._add_paragraph_cn(doc, "图1 泥炭沉积年代-深度关系曲线", font_size=10, bold=True, alignment="center")
        else:
            self._add_paragraph_cn(doc, "年代学数据待补充。", font_size=12)

        self._add_heading_cn(doc, "四、多指标万年演化时序", level=1)
        self._add_paragraph_cn(doc,
            "下图同步展示了碳同位素（δ¹³C）、湿度指数、温度距平和重金属富集指数的万年波动趋势。"
            "各面板共享同一时间轴（yr BP），光标联动显示各指标在同一年代的数值。"
            "浅色填充区域指示不同地层单元，便于识别气候阶段。",
            font_size=12, alignment="justify"
        )

        ts_png = figure_paths.get("timeseries", {}).get("png", "")
        if ts_png:
            self._add_image_safe(doc, ts_png, width_inches=6.2)
            self._add_paragraph_cn(doc, "图2 湿地泥炭多指标万年演化联动时序图", font_size=10, bold=True, alignment="center")

        self._add_heading_cn(doc, "五、多指标相关性分析", level=1)
        self._add_paragraph_cn(doc,
            "采用皮尔逊相关分析量化各代用指标之间的线性关系。"
            "显著性水平设定为α=0.05，*表示p<0.05，**表示p<0.01，***表示p<0.001。",
            font_size=12, alignment="justify"
        )

        corr_png = figure_paths.get("correlation_heatmap", {}).get("png", "")
        if corr_png:
            self._add_image_safe(doc, corr_png, width_inches=5.0)
            self._add_paragraph_cn(doc, "图3 多指标皮尔逊相关性热力图", font_size=10, bold=True, alignment="center")

        significant_corr = corr_report.get("significant_correlations", pd.DataFrame())
        if len(significant_corr) > 0:
            sig_display = significant_corr.copy()
            sig_display.columns = [
                "变量1", "变量2", "Pearson相关系数", "p值", "统计显著", "相关强度"
            ]
            self._add_dataframe_table(doc, sig_display, "表2 统计显著的指标间相关关系 (α=0.05)")

        self._add_heading_cn(doc, "六、主要结论", level=1)

        conclusions = self._generate_conclusions(timeseries_df, corr_report)
        for i, conclusion in enumerate(conclusions, 1):
            self._add_paragraph_cn(doc, f"({i}) {conclusion}", font_size=12, alignment="justify")

        self._add_heading_cn(doc, "附录：数据统计摘要", level=1)

        numeric_cols = timeseries_df.select_dtypes(include=[np.number]).columns
        summary_stats = timeseries_df[numeric_cols].describe().T
        summary_stats = summary_stats[["count", "mean", "std", "min", "25%", "50%", "75%", "max"]]
        summary_stats.index.name = "指标"
        summary_display = summary_stats.reset_index()
        col_name_map = {
            "index": "指标", "count": "样本数", "mean": "均值", "std": "标准差",
            "min": "最小值", "25%": "25%分位", "50%": "中位数", "75%": "75%分位", "max": "最大值"
        }
        summary_display = summary_display.rename(columns=col_name_map)
        self._add_dataframe_table(doc, summary_display, "表A1 各指标统计摘要")

        try:
            doc.save(output_path)
        except PermissionError:
            alt_path = output_path.replace(".docx", "_v2.docx")
            doc.save(alt_path)
            output_path = alt_path

        return output_path

    def _generate_conclusions(self, timeseries_df: pd.DataFrame, corr_report: Dict) -> List[str]:
        """根据数据分析结果生成自动结论"""
        conclusions = []

        if "humidity" in timeseries_df.columns:
            recent_humidity = timeseries_df["humidity"].tail(50).mean()
            early_humidity = timeseries_df["humidity"].head(50).mean()
            if recent_humidity > early_humidity:
                trend = "升高"
            else:
                trend = "降低"
            conclusions.append(
                f"湿度指数显示研究区晚全新世以来整体呈{trend}趋势，"
                f"从早全新世的{early_humidity:.2f}变化至近期的{recent_humidity:.2f}。"
            )

        if "temperature" in timeseries_df.columns:
            temp_std = timeseries_df["temperature"].std()
            temp_max = timeseries_df["temperature"].max()
            temp_min = timeseries_df["temperature"].min()
            conclusions.append(
                f"温度距平在过去一万年间波动幅度达{temp_max - temp_min:.2f}°C，"
                f"标准差为{temp_std:.2f}°C，表明区域气候经历了显著的冷暖交替。"
            )

        if "delta13C" in timeseries_df.columns:
            c13_mean = timeseries_df["delta13C"].mean()
            conclusions.append(
                f"有机碳δ¹³C平均值为{c13_mean:.2f}‰，反映了湿地生态系统碳同位素组成的典型特征。"
            )

        significant = corr_report.get("significant_correlations", pd.DataFrame())
        if len(significant) > 0:
            top = significant.iloc[0]
            direction = "正相关" if top["pearson_r"] > 0 else "负相关"
            conclusions.append(
                f"统计分析表明{top['variable_1']}与{top['variable_2']}呈显著{direction}关系"
                f"(r = {top['pearson_r']:.3f}, p = {top['p_value']:.2e})。"
            )

        if len(conclusions) == 0:
            conclusions.append("数据分析结果待进一步解读。")

        return conclusions

    def generate_cross_region_report(self,
                                    comparison_results: Dict,
                                    figure_paths: Dict[str, Dict[str, str]],
                                    group_id: Optional[str] = None,
                                    filename: Optional[str] = None) -> str:
        """
        生成跨区域泥炭钻孔对比分析Word报告

        Args:
            comparison_results: 对比分析结果字典
            figure_paths: 图表文件路径字典
            group_id: 对比组ID
            filename: 输出文件名（不含扩展名）

        Returns:
            报告文件完整路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请安装python-docx库以生成Word报告：pip install python-docx")

        if group_id is None:
            group_id = "cross_region_comparison"
        if filename is None:
            filename = f"{group_id}_comparison_report"

        doc = Document()
        self._set_default_font(doc)

        title = doc.add_heading("", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("中国湿地泥炭钻孔古环境跨区域对比研究报告")
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = subtitle.add_run("Cross-Region Paleoenvironmental Comparison of Peat Cores")
        sr.font.size = Pt(12)
        sr.italic = True
        sr.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

        self._add_report_header(doc)
        self._add_summary_table_cross(doc, comparison_results)

        group_config = comparison_results.get("comparison_group", {})
        self._add_report_section(
            doc, "1 研究区概况与对比组设置",
            [
                f"本次对比研究采用对比组「{group_config.get('name', group_id)}」，"
                f"共纳入 {len(comparison_results.get('core_timeseries', {}))} 个湿地泥炭钻孔，"
                f"覆盖多个地理气候分区。{group_config.get('description', '')}",
                "钻孔位置、区域分布和海拔高程详见下表。"
            ]
        )
        self._add_core_metadata_table(doc, comparison_results.get("core_metadata"))

        self._add_report_section(
            doc, "2 多钻孔时序联动对比",
            [
                "将各钻孔古环境指标统一校准至相同时间网格（10年分辨率）后，"
                "开展多区域时序联动对比。下图分别展示湿度指数、温度距平和碳同位素的跨区域时空演化。"
            ]
        )

        ts_fig_keys = ["timeseries_comparison_humidity", "timeseries_comparison_temperature",
                      "timeseries_comparison_delta13C"]
        ts_captions = [
            "图2-1 多区域湿地钻孔湿度指数万年波动对比",
            "图2-2 多区域湿地钻孔温度距平万年波动对比",
            "图2-3 多区域湿地钻孔有机碳δ¹³C同位素演化对比"
        ]
        for fk, cap in zip(ts_fig_keys, ts_captions):
            if fk in figure_paths and "png" in figure_paths[fk]:
                self._add_figure_with_caption(doc, figure_paths[fk]["png"], cap)

        self._add_report_section(
            doc, "3 分时段古环境统计对比",
            [
                "按考古气候地层单位（全新世早-中-晚期、末次冰消期）对各钻孔的古环境指标"
                "进行分段统计。区域-时段热力图直观展现不同区域的古环境梯度。"
            ]
        )

        self._add_period_summary_table(doc, comparison_results.get("period_statistics"))

        heatmap_keys = ["region_period_heatmap_humidity", "region_period_heatmap_temperature"]
        heatmap_captions = [
            "图3-1 各区域钻孔分时段湿度指数均值热力图",
            "图3-2 各区域钻孔分时段温度距平热力图"
        ]
        for fk, cap in zip(heatmap_keys, heatmap_captions):
            if fk in figure_paths and "png" in figure_paths[fk]:
                self._add_figure_with_caption(doc, figure_paths[fk]["png"], cap)

        self._add_report_section(
            doc, "4 钻孔间古环境差异与区域聚类",
            [
                "基于各时段古环境指标的欧氏距离构建钻孔间差异矩阵，"
                "采用层次聚类（UPGMA）方法识别区域古环境相似性分组。"
            ]
        )

        diff_keys = ["diff_heatmap_combined", "diff_heatmap_humidity_mean", "diff_heatmap_temperature_mean"]
        diff_captions = [
            "图4-1 多钻孔综合古环境差异热力图（加权欧氏距离）",
            "图4-2 多钻孔湿度指数差异热力图",
            "图4-3 多钻孔温度距平差异热力图"
        ]
        for fk, cap in zip(diff_keys, diff_captions):
            if fk in figure_paths and "png" in figure_paths[fk]:
                self._add_figure_with_caption(doc, figure_paths[fk]["png"], cap)

        clusters = comparison_results.get("clusters")
        if clusters is not None and len(clusters) > 0:
            p = doc.add_paragraph()
            run = p.add_run("表4-1 区域古环境聚类分组结果")
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_dataframe_table(doc, clusters[["core_id", "location", "region", "cluster"]])

        self._add_report_section(
            doc, "5 空间相关性分析",
            [
                "对不同时间节点的钻孔指标与海拔、纬度等空间因子进行皮尔逊相关分析，"
                "探索古环境梯度的空间驱动机制。"
            ]
        )

        spatial_corr = comparison_results.get("spatial_correlation")
        if spatial_corr is not None and len(spatial_corr) > 0:
            self._add_spatial_correlation_summary(doc, spatial_corr)
            if "spatial_correlation" in figure_paths and "png" in figure_paths["spatial_correlation"]:
                self._add_figure_with_caption(
                    doc, figure_paths["spatial_correlation"]["png"],
                    "图5-1 湿度指数与海拔/纬度空间相关性演变"
                )

        self._add_report_section(
            doc, "6 跨区域对比结论",
            self._generate_cross_region_conclusions(comparison_results)
        )

        output_path = self.config.get_report_path(f"{filename}.docx")
        doc.save(output_path)
        return output_path

    def _add_summary_table_cross(self, doc, comparison_results: Dict) -> None:
        """添加跨区域对比摘要表"""
        p = doc.add_paragraph()
        run = p.add_run("摘要")
        run.bold = True
        run.font.size = Pt(14)

        cores = comparison_results.get("core_timeseries", {})
        aligned = comparison_results.get("aligned_data")
        period_stats = comparison_results.get("period_statistics")

        summary_text = (
            f"本次研究共对比分析 {len(cores)} 个湿地泥炭钻孔的多指标古环境记录，"
        )
        if aligned is not None:
            age_min = aligned["age_yrBP"].min()
            age_max = aligned["age_yrBP"].max()
            summary_text += (
                f"覆盖年代范围 {age_min:.0f}–{age_max:.0f} yr BP，"
                f"以 10 年统一分辨率对齐，总计 {len(aligned)} 组观测数据。"
            )
        if period_stats is not None:
            n_periods = period_stats["period_label"].nunique()
            summary_text += f"按 {n_periods} 个考古气候时段进行统计对比分析。"

        doc.add_paragraph(summary_text)

    def _add_core_metadata_table(self, doc, meta_df: Optional[pd.DataFrame]) -> None:
        """添加钻孔元数据表"""
        if meta_df is None:
            doc.add_paragraph("（钻孔元数据暂缺）")
            return

        p = doc.add_paragraph()
        run = p.add_run("表1-1 研究钻孔基本信息")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        display_cols = [c for c in ["core_id", "location", "region",
                                    "latitude", "longitude", "elevation_m", "description"]
                       if c in meta_df.columns]
        self._add_dataframe_table(doc, meta_df[display_cols])

    def _add_period_summary_table(self, doc, period_stats: Optional[pd.DataFrame]) -> None:
        """添加时段统计摘要表"""
        if period_stats is None or len(period_stats) == 0:
            doc.add_paragraph("（分时段统计数据暂缺）")
            return

        p = doc.add_paragraph()
        run = p.add_run("表3-1 各钻孔分时段古环境指标统计摘要")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cols = ["core_id", "location", "period_label", "sample_count"]
        for proxy in ["humidity", "temperature", "delta13C"]:
            if f"{proxy}_mean" in period_stats.columns:
                cols.extend([f"{proxy}_mean", f"{proxy}_std", f"{proxy}_trend"])

        avail_cols = [c for c in cols if c in period_stats.columns]
        region_col = "region" if "region" in period_stats.columns else None
        if region_col and region_col not in avail_cols:
            avail_cols.insert(2, region_col)
        display_df = period_stats[avail_cols].head(30)
        self._add_dataframe_table(doc, display_df)

        if len(period_stats) > 30:
            note = doc.add_paragraph()
            nr = note.add_run(f"（注：仅展示前30行，完整统计共 {len(period_stats)} 条记录）")
            nr.italic = True
            nr.font.size = Pt(9)

    def _add_spatial_correlation_summary(self, doc, spatial_corr: pd.DataFrame) -> None:
        """添加空间相关性摘要"""
        p = doc.add_paragraph()
        run = p.add_run("表5-1 不同时段湿度空间相关性统计")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        summary = spatial_corr.groupby("spatial_variable").agg(
            sample_count=("pearson_r", "count"),
            mean_r=("pearson_r", "mean"),
            max_r=("pearson_r", "max"),
            min_r=("pearson_r", "min"),
            significant_count=("p_value", lambda x: (x < 0.05).sum())
        ).reset_index()

        var_display = {
            "elevation_m": "海拔 (m)",
            "latitude": "纬度 (°N)",
            "longitude": "经度 (°E)"
        }
        summary["空间因子"] = summary["spatial_variable"].map(var_display).fillna(summary["spatial_variable"])
        display = summary[["空间因子", "sample_count", "mean_r", "max_r", "min_r", "significant_count"]]
        display.columns = ["空间因子", "统计时段数", "平均r", "最大r", "最小r", "显著(p<0.05)时段数"]
        display = display.round(3)
        self._add_dataframe_table(doc, display)

    def _generate_cross_region_conclusions(self, comparison_results: Dict) -> List[str]:
        """生成跨区域对比自动结论"""
        conclusions = []

        cores = comparison_results.get("core_timeseries", {})
        n_cores = len(cores)
        if n_cores >= 2:
            conclusions.append(
                f"本研究通过{n_cores}个跨区域泥炭钻孔的多指标对比分析，"
                f"揭示了不同地理单元湿地古环境演化的共性规律与区域差异性。"
            )

        diff_combined = comparison_results.get("difference_matrices", {}).get("combined")
        if diff_combined is not None and len(diff_combined) >= 3:
            values = diff_combined.values.copy()
            np.fill_diagonal(values, np.nan)
            if np.any(~np.isnan(values)):
                avg_diff = np.nanmean(values)
                max_diff = np.nanmax(values)
                max_pos = np.unravel_index(np.nanargmax(values), values.shape)
                most_diff_pair = f"{diff_combined.index[max_pos[0]]} vs {diff_combined.columns[max_pos[1]]}"
                conclusions.append(
                    f"综合古环境差异分析表明，钻孔间平均欧氏距离为{avg_diff:.3f}，"
                    f"其中差异最大的钻孔对为「{most_diff_pair}」(距离={max_diff:.3f})，"
                    f"反映显著的区域古环境梯度。"
                )

        period_stats = comparison_results.get("period_statistics")
        if period_stats is not None and len(period_stats) > 0:
            if "humidity_mean" in period_stats.columns and "region" in period_stats.columns:
                region_means = period_stats.groupby("region")["humidity_mean"].mean().sort_values(ascending=False)
                if len(region_means) >= 2:
                    wettest = region_means.index[0]
                    driest = region_means.index[-1]
                    conclusions.append(
                        f"区域湿度对比显示「{wettest}」区域整体湿度最高（均值={region_means.iloc[0]:.2f}），"
                        f"而「{driest}」区域最为干燥（均值={region_means.iloc[-1]:.2f}），"
                        f"与现代气候格局基本吻合。"
                    )

        clusters = comparison_results.get("clusters")
        if clusters is not None and len(clusters) > 0:
            n_clusters = clusters["cluster"].nunique()
            conclusions.append(
                f"层次聚类将{n_cores}个钻孔划分为{n_clusters}个古环境相似组，"
                f"组内钻孔在湿度、温度演化上表现出同步性特征。"
            )

        spatial_corr = comparison_results.get("spatial_correlation")
        if spatial_corr is not None and len(spatial_corr) > 0:
            sig = spatial_corr[spatial_corr["p_value"] < 0.05]
            if len(sig) > 0:
                top_sig = sig.loc[sig["pearson_r"].abs().idxmax()]
                var_disp = {"elevation_m": "海拔", "latitude": "纬度", "longitude": "经度"}
                var_name = var_disp.get(top_sig["spatial_variable"], top_sig["spatial_variable"])
                direction = "正相关" if top_sig["pearson_r"] > 0 else "负相关"
                conclusions.append(
                    f"空间相关性分析显示在 {top_sig['age_yrBP']:.0f} yr BP 左右，"
                    f"湿度与{var_name}呈显著{direction}(r={top_sig['pearson_r']:.3f}, "
                    f"p={top_sig['p_value']:.2e})，指示古环境梯度的空间驱动机制。"
                )

        if len(conclusions) == 0:
            conclusions.append("跨区域对比分析结果待进一步深入解读。")

        return conclusions

    def _add_dataframe_table(self, doc, df: pd.DataFrame) -> None:
        """将DataFrame添加为Word表格"""
        from docx.shared import Pt
        from docx.oxml.ns import qn

        df = df.copy()
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        df[numeric_cols] = df[numeric_cols].apply(lambda x: x.round(3))
        df = df.fillna("-")

        n_rows, n_cols = df.shape
        table = doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.style = "Light Grid Accent 1"

        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col_name)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for i in range(n_rows):
            for j in range(n_cols):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(df.iloc[i, j])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()
